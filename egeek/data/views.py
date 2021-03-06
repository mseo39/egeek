from django.db import connection
from django.shortcuts import get_object_or_404, render, redirect
from .forms import uploadfile_form
import pandas as pd
from .models import Uploadfile, dorm1_data, dorm2_data,dorm3_data, old_dorm1_data, old_dorm2_data,old_dorm3_data,global_dorm_data, overnight_list,overnight_stay,qrfile
import qrcode
from datetime import datetime
import calendar
from django.contrib.auth.decorators import login_required
import openpyxl
from socket import *
from threading import *

#테스트
@login_required(login_url='/accounts/login/')
def main(request):
    return render(request, 'home.html')

#데이터베이스에 칼럼 가져오기
def dorm_search(dorm, student_number):
    if(dorm=="향1"):
        dorm_= get_object_or_404(dorm1_data, student_number=student_number)
    elif(dorm=="향2"):
        dorm_= get_object_or_404(dorm2_data, student_number=student_number)
    elif(dorm=="향3"):
        dorm_= get_object_or_404(dorm3_data, student_number=student_number)
    elif(dorm=="학성사1"):
        dorm_= get_object_or_404(old_dorm1_data, student_number=student_number)
    elif(dorm=="학성사2"):
        dorm_= get_object_or_404(old_dorm2_data, student_number=student_number)
    elif(dorm=="학성사3"):
        dorm_= get_object_or_404(old_dorm3_data, student_number=student_number)
    elif(dorm=="글로벌빌리지"):
        dorm_= get_object_or_404(global_dorm_data, student_number=student_number)
    return dorm_

def detail(request,dorm, student_number):
    dorm_=dorm_search(dorm, student_number)
    return render(request, 'detail.html', {'dorm_data' : dorm_})

@login_required(login_url='/accounts/login/')
#학생들의 정보가 적혀있는 엑셀파일 업로드
def upload_file(request):
    if request.method=="POST":
        form=Uploadfile()
        file=request.FILES['file']

        title,i=str(file).split('.')
        form.title=title
        form.file=file
        form.save()

    file_form=uploadfile_form()
    return render(request, 'main.html', {'file_form':file_form, 'dis':'disabled'})

import openpyxl
from PIL import Image
from docx import Document
from docx.shared import Inches

#엑셀에 있는 정보를 데이터베이스에 저장
#링크, 기숙사명, 학번 정보를 가지고 있는 qr코드를 생성
def excel_to_db(document,row,file):
    if(row[1][row[1].index[0]]=="향1"):
        dorm=dorm1_data()
    elif(row[1][row[1].index[0]]=="향2"):
        dorm=dorm2_data()
    elif(row[1][row[1].index[0]]=="향3"):
        dorm=dorm3_data()
    elif(row[1][row[1].index[0]]=="학성사1"):
        dorm= old_dorm1_data()
    elif(row[1][row[1].index[0]]=="학성사2"):
        dorm=old_dorm2_data()
    elif(row[1][row[1].index[0]]=="학성사3"):
        dorm=old_dorm3_data()
    elif(row[1][row[1].index[0]]=="글로벌빌리지"):
        dorm=global_dorm_data()
    dorm.dorm=row[1][row[1].index[0]]
    dorm.dorm_number=row[1][row[1].index[1]]
    dorm.student_number=row[1][row[1].index[2]]
    dorm.file_name=file

    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )

    qr.add_data('http://192.168.0.213:8000/')
    qr.add_data('detail/'+row[1][row[1].index[0]]+'/')
    qr.add_data(row[1][row[1].index[2]])
    img = qr.make_image(fill_color="white", back_color="black")
    img.save('media/qr/{}_qr.png'.format(row[1][row[1].index[2]]))
    dorm.qr_image='qr/{}_qr.png'.format(row[1][row[1].index[2]])
    dorm.save()

    p = document.add_paragraph()
    r = p.add_run()
    r.add_text(row[1][row[1].index[0]])
    r.add_text(row[1][row[1].index[1]])
    r.add_text(str(row[1][row[1].index[2]]))
    r.add_picture('media/qr/{}_qr.png'.format(row[1][row[1].index[2]]))


import os
@login_required(login_url='/accounts/login/')
def select_file(request):
    if request.method=="POST":
        chk_file=request.POST.getlist('file[]')
        select=request.POST['submit']

        for file in chk_file:
            file_= get_object_or_404(Uploadfile, file=file)
            if select=="데이터 올리기":
                file_.chk=1
                file_.save()
                directory="media/{}".format(file_.file)
                document = Document()
                qr_file=qrfile()
                
                df=pd.read_excel(directory, header=0)
                #df=pd.read_excel(directory,header=None)

                i=0
                for row in df.iterrows():
                    i=i+1
                    excel_to_db(document,row, file_.file)           

            elif select=="파일 삭제":
                Uploadfile.objects.filter(file=file_.file).delete()
                os.remove("media/{}".format(file_.file))

        return redirect("select_file")

    not_upload_files=Uploadfile.objects.filter(chk=0)
    upload_files=Uploadfile.objects.filter(chk=1)
    return render(request, 'file.html', {'upload_files':upload_files,'not_upload_files':not_upload_files})


def delete(dorm):
    for i in dorm:
        i.delete()

@login_required(login_url='/accounts/login/')
def delete_data(request):
    if request.method=='POST':
        chk_file=request.POST.getlist('file[]')
        for file in chk_file:
            dorm1=dorm1_data.objects.filter(file_name=file)
            delete(dorm1)
            dorm2=dorm2_data.objects.filter(file_name=file)
            delete(dorm2)
            dorm3=dorm3_data.objects.filter(file_name=file)
            delete(dorm3)
            old_dorm1=old_dorm1_data.objects.filter(file_name=file)
            delete(old_dorm1)
            old_dorm2=old_dorm2_data.objects.filter(file_name=file)
            delete(old_dorm2)
            old_dorm3=old_dorm3_data.objects.filter(file_name=file)
            delete(old_dorm3)
            global_dorm=global_dorm_data.objects.filter(file_name=file)
            delete(global_dorm)
            file_= get_object_or_404(Uploadfile, file=file)
            file_.chk=0
            file_.save()
    not_upload_files=Uploadfile.objects.filter(chk=0)
    upload_files=Uploadfile.objects.filter(chk=1)
    return render(request, 'file.html', {'upload_files':upload_files,'not_upload_files':not_upload_files})

from django.views.decorators.csrf import csrf_exempt

#외출
outing={}

def outing_call(dorm, dorm_number, studet_number):
    outing[studet_number]=str(dorm)+","+str(dorm_number)
    
    
def outing_call1():
    return outing
    
@csrf_exempt
def select_out(request, dorm):
    if request.method=='POST':
        student_number=request.POST['student_number']
        select=request.POST['submit']

        dorm_=dorm_search(dorm, student_number)
        lists=overnight_stay.objects.filter(student_number=student_number)
        
        if(dorm=="향1" or dorm=="향2" or dorm=="향3"):
            overnight_stay.objects.filter(dorm=dorm, dorm_number=dorm_.dorm_number)

        days=[]
        c= calendar.TextCalendar(calendar.SUNDAY)

        if(select=="외출"):
            #호출
            outing_call(dorm_.dorm, dorm_.dorm_number,dorm_.student_number)

            return render(request, 'outing.html',{'dorm_data':dorm_})
        #외박form
        elif(select=="감소"):
            month=int(request.POST['month'])-1
            if(month<datetime.today().month):
                month=datetime.today().month
        
        elif(select=="증가"):
            month=int(request.POST['month'])+1
            if(month==13):
                month=12

        else:
            month=datetime.today().month
        
        
        for i in c.itermonthdays(datetime.today().year,month):
            if i==0:
                days.append([" ","disabled"])
            else:
                days.append([i," "])

        already=overnight_stay.objects.filter(month=month,student_number=student_number)
        index=0
        for i in c.itermonthdays(datetime.today().year,month):
            if i!=0:
                break
            index=index+1

        for j in already:
            days[index+j.day-1][1]="disabled"

        return render(request, 'form.html',{'dorm_data':dorm_, "days":days, "month":month})

@csrf_exempt
def delete_date(request, dorm):
    if request.method=='POST':
        student_number=request.POST['student_number']
        select=request.POST['submit']

        dorm_=dorm_search(dorm, student_number)
        
        if(dorm=="향1" or dorm=="향2" or dorm=="향3"):
            overnight_stay.objects.filter(dorm=dorm, dorm_number=dorm_.dorm_number)

        days=[]
        c= calendar.TextCalendar(calendar.SUNDAY)

        if(select=="외출"):
            return render(request, 'outing.html',{'dorm_data':dorm_})
        #외박form
        elif(select=="감소"):
            month=int(request.POST['month'])-1
            if(month<datetime.today().month):
                month=datetime.today().month
        
        elif(select=="증가"):
            month=int(request.POST['month'])+1
            if(month==13):
                month=12

        else:
            month=datetime.today().month
            
        for i in c.itermonthdays(datetime.today().year,month):
            if i==0:
                days.append([" ","disabled"])
            else:
                days.append([i,"disabled"])

        already=overnight_stay.objects.filter(month=month,student_number=student_number)
        
        index=0

        for i in c.itermonthdays(datetime.today().year,month):
            if i!=0:
                break
            index=index+1

        for j in already:
            days[index+j.day-1][1]=" "

        return render(request, 'form_delete.html',{'dorm_data':dorm_, "days":days, "month":month})

from django.http import HttpResponseRedirect,HttpResponse
from django.contrib import messages
from socket import *
from threading import *
#외박 form
#overnight_stay_month, day, dorm, Dorm_number
def overnight(request):
    
    if request.method=="POST":
        chk_date=request.POST.getlist('date[]')
        student_number= request.POST['student_number']
        dorm= request.POST['dorm']
        month= request.POST['month']
        dorm_=dorm_search(dorm, student_number)

        select=request.POST['submit']

        if(len(chk_date)!=0):
            if select=="신청하기":
            
                for i in chk_date:
                    overnight=overnight_stay()
                    overnight.month=month
                    overnight.day=i
                    overnight.dorm=dorm
                    overnight.dorm_number=dorm_.dorm_number
                    overnight.student_number=student_number
                    overnight.save()

            elif select=="취소하기":
                for i in chk_date:
                    overnight_stay.objects.filter(dorm=dorm, dorm_number=dorm_.dorm_number, month=month,day=i).delete()


            count=overnight_stay.objects.filter(dorm=dorm, dorm_number=dorm_.dorm_number, month=month,day=i).count()
            list=overnight_list()
            if (dorm=="향1" or dorm=="향2" or dorm=="향3" or dorm=="학성사2" or dorm=="글로벌빌리지"):
                if count==2:
                    list.month=month
                    list.day=i
                    list.dorm=dorm
                    list.dorm_number=dorm_.dorm_number
                    list.save()

                elif(dorm=="학성사1" or dorm=="학성사3"):
                    if count==4:
                        list.month=month
                        list.day=i
                        list.dorm=dorm
                        list.dorm_number=dorm_.dorm_number
                        list.save()


        else:
              return render(request, 'result1.html', {'error':'날짜를 선택해주세요'})

    return render(request, 'result.html', {'dorm_data':dorm_, 'date':chk_date, 'month': month})


def qr(request):
    return render(request, 'qr.html')

import json
from django.http import JsonResponse

def call(request,option):
    
    #외출 신청이 왔을 때
    if (option=="outing"):
        outing1=outing_call1()
        #외출 데이터를 보내줌
        a_json=json.dumps(outing1)
        print(outing1)
        #보낸데이터는 삭제
        #for key in outing1.items():
         #   del outing[key]

        print("외출")
    
    #외박신청이 왔을 때
    elif(option=="overnight"):
        overnight={}
        list=overnight_list.objects.filter(month=datetime.today().month,day=datetime.today().day)

        for i in list:
            overnight[i.dorm]=i.dorm_number
        #외박 데이터를 보내줌
        a_json=json.dumps(overnight)
        #보낸데이터는 삭제
        list.delete()

        print("외박")
        
    return HttpResponse(a_json)